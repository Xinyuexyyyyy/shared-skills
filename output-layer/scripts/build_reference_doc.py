#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TARGET_REFERENCE = ROOT / "templates" / "docx" / "formal-zh-reference.docx"
DEFAULT_REFERENCE = TARGET_REFERENCE

SIZE_MAP = {
    "小二": 18.0,
    "三号": 16.0,
    "小五": 9.0,
    "小三": 15.0,
    "四号": 14.0,
    "小四": 12.0,
    "五号": 10.5,
}


def _set_run_font(style, east_asia: str, latin: str, size_pt: float, bold: bool = False) -> None:
    style.font.name = latin
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size_pt)
    style.font.bold = bold


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build reference.docx for output-layer")
    parser.add_argument("--body-font", default="宋体")
    parser.add_argument("--body-size", default="小四")
    parser.add_argument("--title-font", default="黑体")
    parser.add_argument("--title-size", default="小二")
    parser.add_argument("--title-weight", default="bold")
    parser.add_argument("--title-align", default="center")
    parser.add_argument("--heading-font", default="黑体")
    parser.add_argument("--heading-weight", default="bold")
    parser.add_argument("--heading1-size", default="16")
    parser.add_argument("--heading2-size", default="14")
    parser.add_argument("--heading3-size", default="12")
    parser.add_argument("--line-spacing", default="1.5")
    parser.add_argument("--first-line-indent", default="2char")
    parser.add_argument("--spacing-before", default="0")
    parser.add_argument("--spacing-after", default="0")
    parser.add_argument("--alignment", default="justify")
    parser.add_argument("--heading1-spacing-before", default="12")
    parser.add_argument("--heading1-spacing-after", default="6")
    parser.add_argument("--heading2-spacing-before", default="10")
    parser.add_argument("--heading2-spacing-after", default="6")
    parser.add_argument("--heading3-spacing-before", default="8")
    parser.add_argument("--heading3-spacing-after", default="3")
    parser.add_argument("--heading-align", default="left")
    parser.add_argument("--margin", default="normal")
    parser.add_argument("--top-margin", default="")
    parser.add_argument("--bottom-margin", default="")
    parser.add_argument("--left-margin", default="")
    parser.add_argument("--right-margin", default="")
    parser.add_argument("--page-number", default="off")
    parser.add_argument("--page-number-align", default="center")
    parser.add_argument("--first-page-number", default="off")
    parser.add_argument("--header-text", default="")
    parser.add_argument("--footer-text", default="")
    parser.add_argument("--header-align", default="center")
    parser.add_argument("--footer-align", default="center")
    parser.add_argument("--header-font", default="宋体")
    parser.add_argument("--header-size", default="小五")
    parser.add_argument("--footer-font", default="宋体")
    parser.add_argument("--footer-size", default="小五")
    parser.add_argument("--page-number-font", default="宋体")
    parser.add_argument("--page-number-size", default="小五")
    parser.add_argument("--base-reference", default=str(DEFAULT_REFERENCE))
    parser.add_argument("--output", default=str(TARGET_REFERENCE))
    return parser.parse_args()


def _resolve_size(label: str) -> float:
    try:
        return float(label)
    except ValueError:
        return SIZE_MAP.get(label, 12.0)


def _resolve_alignment(name: str):
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    return mapping.get(name, WD_ALIGN_PARAGRAPH.LEFT)


def _apply_line_spacing(paragraph_format, value: str) -> None:
    try:
        numeric = float(value)
    except ValueError:
        numeric = 1.5
    if numeric >= 10:
        paragraph_format.line_spacing = Pt(numeric)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    else:
        paragraph_format.line_spacing = numeric
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def _apply_section_margins(section, margin: str) -> None:
    if margin == "narrow":
        section.top_margin = Cm(1.91)
        section.bottom_margin = Cm(1.91)
        section.left_margin = Cm(1.91)
        section.right_margin = Cm(1.91)
    elif margin == "wide":
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(4.0)
        section.right_margin = Cm(4.0)
    else:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def _maybe_override_margin(section, override_name: str, value: str | None) -> None:
    if not value:
        return
    try:
        setattr(section, override_name, Cm(float(value)))
    except ValueError:
        return


def _style_run(run, east_asia: str, size_label: str) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(_resolve_size(size_label))


def _set_para_run_font(style, east_asia: str, size_pt: float, bold: bool = False) -> None:
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    style.font.size = Pt(size_pt)
    style.font.bold = bold


def _ensure_paragraph_style(doc: Document, name: str, base: str = "Normal"):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _add_page_number(section, align: str, font_name: str, size_label: str) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = _resolve_alignment(align)
    run = paragraph.add_run()
    _style_run(run, font_name, size_label)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _set_header_text(section, text: str, align: str, font_name: str, size_label: str) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = _resolve_alignment(align)
    if text:
        run = paragraph.add_run(text)
        _style_run(run, font_name, size_label)


def _set_footer_text(section, text: str, align: str, font_name: str, size_label: str) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = _resolve_alignment(align)
    if text:
        run = paragraph.add_run(text)
        _style_run(run, font_name, size_label)


def _set_title_style(doc: Document, font_name: str, size_label: str, weight: str, align: str) -> None:
    title_style = doc.styles["Title"]
    _set_para_run_font(title_style, font_name, _resolve_size(size_label), bold=weight == "bold")
    title_style.paragraph_format.alignment = _resolve_alignment(align)


def _style_block_text(doc: Document, font_name: str, size_label: str) -> None:
    style = _ensure_paragraph_style(doc, "Block Text")
    _set_para_run_font(style, font_name, _resolve_size(size_label))
    style.paragraph_format.left_indent = Cm(0)
    style.paragraph_format.right_indent = Cm(0)
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _style_block_quote(doc: Document, font_name: str, size_label: str) -> None:
    for style_name in ["Block Quote", "Quote"]:
        style = _ensure_paragraph_style(doc, style_name)
        _set_para_run_font(style, font_name, _resolve_size(size_label))
        style.paragraph_format.left_indent = Cm(0.74)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _style_list_paragraph(doc: Document, font_name: str, size_label: str) -> None:
    style = _ensure_paragraph_style(doc, "List Paragraph")
    _set_para_run_font(style, font_name, _resolve_size(size_label))
    style.paragraph_format.left_indent = Cm(0.74)
    style.paragraph_format.first_line_indent = Cm(-0.74)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _style_list_like_styles(doc: Document, font_name: str, size_label: str) -> None:
    for style_name in ["List Paragraph", "Compact"]:
        style = _ensure_paragraph_style(doc, style_name)
        _set_para_run_font(style, font_name, _resolve_size(size_label))
        if style_name == "Compact":
            style.paragraph_format.left_indent = Cm(0.55)
            style.paragraph_format.first_line_indent = Cm(-0.55)
        else:
            style.paragraph_format.left_indent = Cm(0.74)
            style.paragraph_format.first_line_indent = Cm(-0.74)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _style_body_text(doc: Document, font_name: str, size_label: str) -> None:
    for style_name in ["Body Text", "BodyText", "First Paragraph", "FirstParagraph"]:
        style = _ensure_paragraph_style(doc, style_name)
        _set_para_run_font(style, font_name, _resolve_size(size_label))
        style.paragraph_format.left_indent = Cm(0)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _set_paragraph_shading(paragraph_style, fill: str) -> None:
    p_pr = paragraph_style._element.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _set_paragraph_borders(paragraph_style, color: str = "D9D9D9", size: str = "6") -> None:
    p_pr = paragraph_style._element.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), color)


def _style_source_code(doc: Document, font_name: str, size_label: str) -> None:
    style = _ensure_paragraph_style(doc, "Source Code")
    _set_para_run_font(style, font_name, _resolve_size(size_label))
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.paragraph_format.left_indent = Cm(0.55)
    style.paragraph_format.right_indent = Cm(0.55)
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.space_before = Pt(8)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_shading(style, "F5F5F5")
    _set_paragraph_borders(style)
    try:
        verbatim = doc.styles["Verbatim Char"]
        verbatim.font.name = "Times New Roman"
        verbatim._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        verbatim._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        verbatim._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except KeyError:
        pass


def _style_existing_paragraph_styles(doc: Document, alignment) -> None:
    for style_name in ["Body Text", "BodyText", "Compact", "First Paragraph", "FirstParagraph"]:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.paragraph_format.alignment = alignment


def _style_table_cells(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.paragraph_format.left_indent = Cm(0)
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _style_table(doc: Document, font_name: str, size_label: str) -> None:
    style = doc.styles["Table"]
    _set_para_run_font(style, font_name, _resolve_size(size_label))


def main() -> int:
    args = parse_args()
    base_reference = Path(args.base_reference)
    output = Path(args.output)
    if not base_reference.exists():
        raise SystemExit(f"default reference not found: {base_reference}")

    doc = Document(str(base_reference))

    _set_title_style(doc, args.title_font, args.title_size, args.title_weight, args.title_align)
    _style_block_text(doc, args.body_font, args.body_size)
    _style_block_quote(doc, args.body_font, args.body_size)
    _style_list_paragraph(doc, args.body_font, args.body_size)
    _style_list_like_styles(doc, args.body_font, args.body_size)
    _style_body_text(doc, args.body_font, args.body_size)
    _style_source_code(doc, args.body_font, "小五")
    _style_table(doc, args.body_font, args.body_size)
    _style_existing_paragraph_styles(doc, WD_ALIGN_PARAGRAPH.LEFT)

    for section in doc.sections:
        _apply_section_margins(section, args.margin)
        _maybe_override_margin(section, "top_margin", getattr(args, "top_margin", None))
        _maybe_override_margin(section, "bottom_margin", getattr(args, "bottom_margin", None))
        _maybe_override_margin(section, "left_margin", getattr(args, "left_margin", None))
        _maybe_override_margin(section, "right_margin", getattr(args, "right_margin", None))
        if args.header_text:
            _set_header_text(section, args.header_text, args.header_align, args.header_font, args.header_size)
        if args.footer_text:
            _set_footer_text(section, args.footer_text, args.footer_align, args.footer_font, args.footer_size)
        if args.page_number == "on":
            if args.first_page_number == "off":
                section.different_first_page_header_footer = True
                _add_page_number(section, args.page_number_align, args.page_number_font, args.page_number_size)
            else:
                _add_page_number(section, args.page_number_align, args.page_number_font, args.page_number_size)

    normal = doc.styles["Normal"]
    _set_para_run_font(normal, args.body_font, _resolve_size(args.body_size))
    if args.first_line_indent == "2char":
        normal.paragraph_format.first_line_indent = Cm(0.74)
    elif args.first_line_indent == "1char":
        normal.paragraph_format.first_line_indent = Cm(0.37)
    elif args.first_line_indent == "hanging":
        normal.paragraph_format.first_line_indent = Cm(-0.74)
    else:
        normal.paragraph_format.first_line_indent = Cm(0)
    _apply_line_spacing(normal.paragraph_format, args.line_spacing)
    normal.paragraph_format.space_before = Pt(float(args.spacing_before))
    normal.paragraph_format.space_after = Pt(float(args.spacing_after))
    normal.paragraph_format.alignment = _resolve_alignment(args.alignment)

    _style_body_text(doc, args.body_font, args.body_size)
    _style_list_like_styles(doc, args.body_font, args.body_size)

    heading1 = doc.styles["Heading 1"]
    _set_para_run_font(heading1, args.heading_font, _resolve_size(args.heading1_size), bold=args.heading_weight == "bold")
    heading1.paragraph_format.first_line_indent = Cm(0)
    heading1.paragraph_format.left_indent = Cm(0)
    heading1.paragraph_format.space_before = Pt(float(args.heading1_spacing_before))
    heading1.paragraph_format.space_after = Pt(float(args.heading1_spacing_after))
    heading1.paragraph_format.alignment = _resolve_alignment(args.heading_align)

    heading2 = doc.styles["Heading 2"]
    _set_para_run_font(heading2, args.heading_font, _resolve_size(args.heading2_size), bold=args.heading_weight == "bold")
    heading2.paragraph_format.first_line_indent = Cm(0)
    heading2.paragraph_format.left_indent = Cm(0)
    heading2.paragraph_format.space_before = Pt(float(args.heading2_spacing_before))
    heading2.paragraph_format.space_after = Pt(float(args.heading2_spacing_after))
    heading2.paragraph_format.alignment = _resolve_alignment(args.heading_align)

    heading3 = doc.styles["Heading 3"]
    _set_para_run_font(heading3, args.heading_font, _resolve_size(args.heading3_size), bold=args.heading_weight == "bold")
    heading3.paragraph_format.first_line_indent = Cm(0)
    heading3.paragraph_format.left_indent = Cm(0)
    heading3.paragraph_format.space_before = Pt(float(args.heading3_spacing_before))
    heading3.paragraph_format.space_after = Pt(float(args.heading3_spacing_after))
    heading3.paragraph_format.alignment = _resolve_alignment(args.heading_align)

    _style_table_cells(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
