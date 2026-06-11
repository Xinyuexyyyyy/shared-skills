from __future__ import annotations

import subprocess
import shutil
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

try:
    from docx import Document
    from docx.shared import Cm
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    Cm = None
    qn = None


TEXTUAL_STYLES = {
    "Body Text",
    "BodyText",
    "Compact",
    "Block Text",
    "First Paragraph",
    "FirstParagraph",
}

FIRST_LINE_EXEMPT_STYLES = {
    "Title",
    "Body Text",
    "BodyText",
    "Compact",
    "Block Text",
    "Quote",
    "First Paragraph",
    "FirstParagraph",
    "Source Code",
    "List Paragraph",
    "Heading 1",
    "Heading 2",
    "Heading 3",
}

PDF_ENGINE_CANDIDATES = [
    ("weasyprint", "html-styled"),
    ("wkhtmltopdf", "html"),
    ("xelatex", "latex"),
    ("lualatex", "latex"),
    ("pdflatex", "latex"),
]

PDF_STYLE_PATH = Path(__file__).resolve().parents[1] / "templates" / "pdf" / "formal-zh.css"

LUA_FILTER_DIR = Path(__file__).resolve().parents[1] / "lua"


def _set_run_east_asia_font(run, font_name: str) -> None:
    if qn is None:
        return
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _normalize_docx(doc) -> None:
    if Cm is None:
        return
    if doc.paragraphs:
        first = doc.paragraphs[0]
        if first.style.name != "Title":
            first.style = doc.styles["Title"]

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name
        if not paragraph.text.strip():
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.paragraph_format.right_indent = Cm(0)
        if style_name in FIRST_LINE_EXEMPT_STYLES or paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            paragraph.paragraph_format.first_line_indent = Cm(0)
        if style_name in TEXTUAL_STYLES:
            for run in paragraph.runs:
                if run.style and run.style.name == "Verbatim Char":
                    continue
                run.bold = False
        if style_name == "Source Code":
            for run in paragraph.runs:
                run.bold = False
                _set_run_east_asia_font(run, "仿宋GB2312")
                if run.text.strip():
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋GB2312")
    try:
        verbatim = doc.styles["Verbatim Char"]
        verbatim.font.name = "Times New Roman"
        verbatim._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        verbatim._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        verbatim._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋GB2312")
    except KeyError:
        pass


def _normalize_docx_xml(output_path: Path) -> None:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(output_path, "r") as src:
        files = {name: src.read(name) for name in src.namelist()}

    root = ET.fromstring(files["word/document.xml"])
    first_nonempty = None
    for paragraph in root.findall(".//w:body/w:p", ns):
        text = "".join(node.text or "" for node in paragraph.findall(".//w:t", ns)).strip()
        if text:
            first_nonempty = paragraph
            break
    if first_nonempty is None:
        return

    ppr = first_nonempty.find("w:pPr", ns)
    if ppr is None:
        ppr = ET.Element(f"{{{ns['w']}}}pPr")
        first_nonempty.insert(0, ppr)
    pstyle = ppr.find("w:pStyle", ns)
    if pstyle is None:
        pstyle = ET.Element(f"{{{ns['w']}}}pStyle")
        ppr.insert(0, pstyle)
    pstyle.set(f"{{{ns['w']}}}val", "Title")

    files["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with ZipFile(output_path, "w") as dst:
        for name, content in files.items():
            dst.writestr(name, content)


def _build_pdf_html(input_markdown: Path) -> str:
    command = [
        "pandoc",
        str(input_markdown),
        "-f",
        "markdown",
        "-s",
        "-t",
        "html5",
        "--no-highlight",
        "--metadata",
        "pagetitle=output-layer",
    ]
    try:
        result = subprocess.run(command, check=True, capture_output=True, text=True)
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(exc.stderr.strip() or str(exc)) from exc
    return result.stdout


def _inject_pdf_stylesheet(html: str) -> str:
    # Pandoc standalone HTML includes broad browser defaults. The PDF path owns
    # typography itself, so strip that block before adding the formal stylesheet.
    start = html.find("<style>")
    end = html.find("</style>", start)
    if start != -1 and end != -1:
        html = html[:start] + html[end + len("</style>"):]
    if not PDF_STYLE_PATH.exists():
        return html
    link = f'<link rel="stylesheet" href="{PDF_STYLE_PATH.as_uri()}">'
    if "</head>" in html:
        return html.replace("</head>", f"{link}\n</head>", 1)
    return f"<!doctype html><html><head>{link}</head><body>{html}</body></html>"


def render_docx(input_markdown: Path, run_dir: Path, reference_doc: Path | None = None) -> tuple[Path | None, str | None]:
    if not shutil.which("pandoc"):
        return None, "pandoc not installed"

    output_path = run_dir / "output.docx"
    command = [
        "pandoc",
        str(input_markdown),
        "-f",
        "markdown",
        "-t",
        "docx",
        "--no-highlight",
        "-o",
        str(output_path),
    ]
    if reference_doc and reference_doc.exists():
        command.extend(["--reference-doc", str(reference_doc)])

    # Load lua filters to fix known pandoc docx issues
    for lua_file in [
        "image-title-to-caption.lua",
        "preserve-font-color.lua",
        "inline-html-tags.lua",
    ]:
        lua_path = LUA_FILTER_DIR / lua_file
        if lua_path.exists():
            command.extend(["--lua-filter", str(lua_path)])

    try:
        subprocess.run(command, check=True, capture_output=True, text=True)
    except subprocess.CalledProcessError as exc:
        return None, exc.stderr.strip() or str(exc)
    try:
        if Document is not None:
            doc = Document(str(output_path))
            _normalize_docx(doc)
            doc.save(str(output_path))
        else:
            _normalize_docx_xml(output_path)
    except Exception:
        pass
    return output_path, None


def _resolve_pdf_engine() -> tuple[str, str] | None:
    for engine_name, engine_kind in PDF_ENGINE_CANDIDATES:
        if shutil.which(engine_name):
            return engine_name, engine_kind
    return None


def render_pdf(input_markdown: Path, run_dir: Path) -> tuple[Path | None, str | None]:
    if not shutil.which("pandoc"):
        return None, "pandoc not installed"
    resolved = _resolve_pdf_engine()
    if resolved is None:
        engine_list = ", ".join(name for name, _kind in PDF_ENGINE_CANDIDATES)
        return None, f"no pdf engine found; install one of: {engine_list}"

    engine_name, engine_kind = resolved
    output_path = run_dir / "output.pdf"
    if engine_name == "weasyprint" and engine_kind == "html-styled":
        try:
            html = _build_pdf_html(input_markdown)
            html_path = run_dir / "output.pdf.html"
            html_path.write_text(_inject_pdf_stylesheet(html), encoding="utf-8")
            subprocess.run(
                [engine_name, str(html_path), str(output_path)],
                check=True,
                capture_output=True,
                text=True,
            )
        except RuntimeError as exc:
            return None, str(exc)
        except subprocess.CalledProcessError as exc:
            return None, exc.stderr.strip() or str(exc)
        except Exception as exc:
            return None, str(exc)
        return output_path, None

    command = [
        "pandoc",
        str(input_markdown),
        "-f",
        "markdown",
        "-s",
        "-o",
        str(output_path),
        "--pdf-engine",
        engine_name,
    ]
    if engine_kind == "html":
        command.extend(["-t", "html5"])
    try:
        subprocess.run(command, check=True, capture_output=True, text=True)
    except subprocess.CalledProcessError as exc:
        return None, exc.stderr.strip() or str(exc)
    return output_path, None
